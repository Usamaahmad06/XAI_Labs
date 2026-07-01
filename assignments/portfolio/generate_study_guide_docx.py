"""
Generate Word (.docx) from the study guide Markdown.
Usage: python assignments/portfolio/generate_study_guide_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError as exc:
    raise SystemExit("Install: pip install python-docx") from exc

ROOT = Path(__file__).resolve().parents[2]
MD_PATH = Path(__file__).resolve().parent / "XAI_Complete_Study_Guide.md"
OUT_PATH = Path(__file__).resolve().parent / "XAI_Complete_Study_Guide.docx"


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    text = text.strip()
    if not text:
        return
    # Strip simple markdown bold/italic/code
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(6)


def main() -> None:
    if not MD_PATH.is_file():
        raise SystemExit(f"Missing: {MD_PATH}")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("XAI Labs — Complete Study Guide", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(
        "Explainable AI course portfolio — concepts, code, results, alternatives, and Q&A.",
        style="Subtitle",
    )

    in_code = False
    code_lines: list[str] = []

    for raw_line in MD_PATH.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                block = "\n".join(code_lines)
                p = doc.add_paragraph(block)
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("- ") or line.startswith("* "):
            add_paragraph(doc, line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            add_paragraph(doc, re.sub(r"^\d+\.\s", "", line), style="List Number")
        elif line.startswith("---"):
            doc.add_paragraph("")
        elif line.startswith("|"):
            add_paragraph(doc, line)
        else:
            add_paragraph(doc, line)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
